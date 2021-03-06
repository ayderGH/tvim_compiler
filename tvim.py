import argparse
import yaml
import shutil
import os
import subprocess
import re
import logging
import PyPDF2
from docx import Document
import json
from transliterate import translit


logger = logging.getLogger('tvim')
logger.setLevel(logging.INFO)


def get_text_between_braces(text, open_pos=0):
    """
    Получить текст между фигурными скобками.
    """
    counter = 1
    i = open_pos
    while counter > 0 and i < len(text):
        if text[i] == '{':
            counter += 1
        elif text[i] == '}':
            counter -= 1
        if counter == 0:
            return text[open_pos:i], i
        i += 1


class ArticleBase:

    def __init__(self, path):
        self._path = path
        self.text = self.get_text()
        self.article_text = ''

    path = property(lambda self: self._path)
    art_path = property(lambda self: os.path.join(self.path, '__article.tex'))

    def compile(self):
        pass

    def get_text(self):
        tex_file = [f for f in os.listdir(self.path) if f.endswith('.tex')]
        if tex_file:
            tex_file = os.path.join(self.path, tex_file[0])
            with open(tex_file, 'rt') as f:
                text = f.read()
            return text


class VerbatimArticle(ArticleBase):

    def compile(self):
        self.article_text = \
            r'\newpage' + '\n' \
            + self.text
        with open(self.art_path, 'wt') as f:
            f.write(self.article_text)


class Article(ArticleBase):
    """
    Объектная модель статьи журнала.
    """
    def __init__(self, path, lang='ru'):
        super().__init__(path)
        self.lang = lang
        self.title = {}
        self.authors = {}
        self.author_details = []
        self.udc = None
        self.msc2010 = None
        self.abstracts = {}
        self.sections = {}
        self.bibliography = {}
        self.article_text = None
        self.authors_en = None
        self.title_en = None
        self.keywords = {}

    @property
    def id(self):
        s = self.authors_str
        s = s.replace('\\;', '_')
        s = re.sub(' +', '_', s)
        s = re.sub(r'[.,]', '', s)
        try:
            return translit(s, reversed=True)
        except Exception as e:
            return s

    @staticmethod
    def normalize_text(text):
        """
        Удалить:
            textbf
            textit
        """
        text = re.sub(r'\\textbf|\\textit|\\it|\\bf', '', text)
        if text.startswith('{') and text.endswith('}'):
            text = text[1:-1]
        # text = re.sub(r'[{|}]', '', text)
        text = re.sub(r'\s{2,}', ' ', text)
        return text.strip()

    @staticmethod
    def _check_balance_of_parantheses(s):
        counter = 0
        for ch in s:
            if ch == '{':
                counter += 1
            elif ch == '}':
                counter -= 1
        return counter

    def select_tag(self, tag, text, default=None):
        m = re.search(tag + r'{.*?}+', text, flags=re.DOTALL)
        if m:
            t = text[m.start() + len(tag):m.end() - 1]
            # if not self._check_balance_of_parantheses(t):
            #     i = m.end()
            t = re.sub(r'\n', ' ', t, flags=re.DOTALL)
            t = re.sub(r' {2,}', ' ', t, flags=re.DOTALL)
            return t
        else:
            return default

    def extract_title(self):
        """
        Извлечь заголовок.
        """
        text = self.text
        self.title = {}
        title = self.select_tag(r'\\title', text)
        if title:
            title = self.select_tag(r'\\uppercase', title, title)
            self.title['ru'] = title
        else:
            logger.error('Не найден заголовок в {}!'.format(self.path))

    def extract_authors(self):
        """
        Извлечь авторов.
        """
        text = self.text
        # russian case
        self.authors = []
        for m1 in re.finditer(r'\\author{(.*)}', text):
            authors = m1[1]
            authors = re.sub(r'\\[;,.:]+', ' ', authors)
            authors = re.sub(r'\s{2,}', ' ', authors)
            name_abbr_pattern = r'(?P<name>[A-ZА-ЯЁa-zа-яё]{1,2}\.)'
            patronymic_abbr_pattern = r'(?P<patronymic>[A-ZА-ЯЁa-zа-яё]{1,2}\.)'
            family_pattern = r'(?P<family>[A-ZА-Яa-zа-яё]{2,})'
            patterns = [r'{}\s*{}\s*{}'.format(name_abbr_pattern,
                                               patronymic_abbr_pattern,
                                               family_pattern),
                        r'{}\s*{}\s*{}'.format(family_pattern,
                                               name_abbr_pattern,
                                               patronymic_abbr_pattern)]

            for pattern in patterns:
                for m in re.finditer(pattern, authors):
                    self.authors.append({'family': m['family'],
                                         'name': m['name'],
                                         'patronymic': m['patronymic']})
        if not self.authors:
            logger.error('Не найдены авторы в {}!'.format(self.path))

    def extract_en_title_and_authors(self):
        # english case
        text = self.text
        pattern = r'\\begin{abstractX}{(?P<title>.*?)}{(?P<authors>.*?)}'
        m = re.search(pattern, text, flags=re.DOTALL)
        if m:
            self.authors_en = m['authors']
            self.title_en = m['title']
        else:
            logger.error('Не найдены авторы (английский вариант) в {}!'.
                         format(self.path))

    def extract_ru_abstracts(self):
        """
        Извлечь русскую аннотацию.
        """
        text = self.text
        self.abstracts['ru'] = ''
        m = re.search(r'\\begin{abstractXr}\n*'
                      r'{(.*?)}\n*{(.*?)}\n*(.*?)\n*'
                      r'\\end{abstractXr}',
                      text, flags=re.DOTALL)
        if m:
            self.abstracts['ru'] = m[3]
        if not self.abstracts['ru']:
            logger.error('Не найдена русская аннотация в {}!'.format(self.path))

    def extract_en_abstracts(self):
        """
        Извлечь английскую аннотацию.
        """
        self.abstracts['en'] = ''
        text = self.text
        begin_pattern = r'\\begin\{abstractX\}.*'
        m = re.search(begin_pattern, text)
        if m:
            p0 = m.end()
            end_pattern = r'\\end\{abstractX\}'
            m = re.search(end_pattern, text)
            if m:
                p1 = m.start()
                abstract = text[p0:p1]
                self.abstracts['en'] = abstract
        if not self.abstracts['en']:
            logger.error('Не найдена английская аннотация в {}!'.
                         format(self.path))

    def extract_keywords(self):
        self.keywords = {'ru': '', 'en': ''}
        m = re.search(r'\\keywordsr{(.*?)}+', self.text, flags=re.DOTALL)
        if m:
            self.keywords['ru'] = m[1]
        else:
            logger.warning("Не найдены ключевые слова "
                           "на русском языке в {}".format(self.path))

        m = re.search(r'\\keywords{(.*?)}+', self.text, flags=re.DOTALL)
        if m:
            self.keywords['en'] = m[1]
        else:
            logger.warning("Не найдены ключевые слова "
                           "на английском языке в {}".format(self.path))

    def extract_sections(self):
        """
        Извлечь разделы (главы) статьи.
        """
        text = self.text
        pattern = r'\\section\*?\{'
        self.sections = []
        for m in re.finditer(pattern, text):
            t = get_text_between_braces(text, m.end())
            if t:
                self.sections.append(self.normalize_text(t[0]))
        print(self.sections)
        if not self.sections:
            logger.error('Не найдены разделы в {}!'.format(self.path))

    def extract_bibliography(self):
        """
        Извлечь список литературы.
        """
        text = self.text

        m = re.search(r'\\begin{thebibliography}', text)
        if m:
            bib_begin_pos = m.start()
        else:
            bib_begin_pos = None

        m = re.search(r'\\end{thebibliography}', text)
        if m:
            bib_end_pos = m.end()
        else:
            bib_end_pos = None

        if not (bib_begin_pos and bib_end_pos):
            logger.error('There is no the bibliography')
        else:
            bib_text = text[bib_begin_pos:bib_end_pos]
            bib_text = re.sub(r'%\s*(.*)', '', bib_text)
            bib_text = bib_text.replace('\n', ' ')

            bib_ids = []
            for m in re.finditer(r'\\bibitem{(?P<bibitem>.*?)}', bib_text):
                bib_ids.append(m['bibitem'])

            self.bibliography = {}
            for bid in bib_ids:
                p = '\\\\bibitem{{{}}}(?P<text>.*?)(?=\\\\bibitem|\\\\end)'.\
                    format(bid)
                m = re.search(p, bib_text, re.MULTILINE)
                if m:
                    bibtext = m['text']

                    self.bibliography[bid] = bibtext

    def extract_udc(self):
        """
        Извлечь УДК.
        """
        p = r'(?<=УДК\:)\s*(.*)(?=\})'
        m = re.search(p, self.text)
        if m:
            self.udc = m[1]
        else:
            self.udc = '???'
            logger.error('Не найден УДК в {}!'.format(self.path))

    def extract_msc2010(self):
        """
        Извлечь MSC2010.
        """
        p = r'(?<=MSC2010\:)\s*(.*)(?=\})'
        m = re.search(p, self.text)
        if m:
            self.msc2010 = m[1]
        else:
            self.msc2010 = '???'
            logger.error('Не найден MSC2010 в {}!'.format(self.path))

    def update_image_path(self):
        """
        Обновить пути к файлам изображений.
        """
        text = self.article_text
        pos = []
        for m in re.finditer(r'\\includegraphics.*?{(.+?)}', text):
            graph_text = text[m.start():m.end()]
            m_image_name = re.search(r'{(.*)}', graph_text)
            pos.append((m.start() + m_image_name.start(),
                        m.start() + m_image_name.end()))
        d = len(self.path) + 1
        for i, p in enumerate(pos):
            image_name = text[p[0] + i*d+1:p[1] + i*d-1]
            text = text[:p[0] + i*d] + \
                '{{{}/{}}}'.format(self.path, image_name) + \
                text[p[1] + i*d:]
        self.article_text = text

    @property
    def authors_str(self):
        return ', '.join(['{}\\;{}\\;{}'.format(a['family'], a['name'],
                                                a['patronymic'])
                          for a in self.authors])

    @property
    def authors_str_reverse(self):
        return ', '.join(['{}.\\;{}.\\;{}'.format(a['name'][0],
                                                  a['patronymic'][0],
                                                  a['family'])
                          for a in self.authors])

    @property
    def begin_label(self):
        return f'{self.id}_begin'

    @property
    def end_label(self):
        return f'{self.id}_end'

    def add_content_lines(self):
        title = self.title['ru']
        title = re.sub(r'\\footnote{.*?}', '', title)
        ru_con = '\\addcontentsline{{toc}}{{art}}' \
                 '{{\\textbf{{{authors}}} {title}}}'.\
                 format(
                    authors=self.authors_str,
                    title=title)

        en_con = '\\addcontentsline{{tec}}{{art}}' \
                 '{{\\textbf{{{authors}}} {title}}}'.\
                 format(authors=self.authors_en, title=self.title_en)
        return '{}\n{}\n'.format(ru_con, en_con)

    def extract_author_details(self):
        p = r'\\authorInfo\{.*?\}\n{2,}'
        self.author_details = []
        for m in re.finditer(p, self.text, flags=re.DOTALL):
            self.author_details.append(self.text[m.start():m.end()])

    def parse(self):
        self.extract_title()
        self.extract_authors()
        self.extract_author_details()
        self.extract_ru_abstracts()
        self.extract_en_abstracts()
        self.extract_keywords()
        self.extract_en_title_and_authors()
        self.extract_udc()
        self.extract_msc2010()
        self.extract_sections()
        self.extract_bibliography()

    def update_title(self):
        # find \footnote
        p = self.title['ru'].find(r'\footnote')
        if p >= 0:
            main_title = self.title['ru'][:p]
            footnote = self.title['ru'][p:]
            upper_title = '{}{}'.format(main_title.upper(), footnote)
        else:
            upper_title = self.title['ru'].upper()

        if self.title['ru']:
            m = re.search(r'\\title{.*?}+', self.article_text, flags=re.DOTALL)
            if m:
                self.article_text = re.sub(
                    r'\\title{.*?}+', '', self.article_text, flags=re.DOTALL)
                self.article_text = self.article_text[:m.start()] \
                    + '\\title{{{}}}'.format(upper_title) \
                    + self.article_text[m.start():]

    def update_section(self, section):
        # TODO: replace by regular expression
        pattern = f'\\section{{{section}}}'
        bold_pattern = f'\\section{{\\textbf{{{section}}}}}'
        pattern_with_asterix = f'\\section*{{{section}}}'
        bold_pattern_with_asterix = f'\\section*{{\\textbf{{{section}}}}}'

        if pattern in self.article_text:
            self.article_text = self.article_text.replace(pattern,
                                                          bold_pattern)
        elif pattern_with_asterix in self.article_text:
            self.article_text = self.article_text.replace(
                pattern_with_asterix, bold_pattern_with_asterix)

        if bold_pattern not in self.article_text \
                and bold_pattern_with_asterix not in self.article_text:
            logger.warning(f"Раздел `{section}` "
                           f"в статье `{self.title['ru']}` "
                           f"не был выделен жирным!")

    art_path = property(lambda self: os.path.join(self.path, '__article.tex'))

    def remove_russian_abstract(self):
        pattern = r'\\begin{abstractXr}\n*' \
                  r'{(.*?)}\n*{(.*?)}\n*(.*?)\n*' \
                  r'\\end{abstractXr}'
        self.article_text = re.sub(pattern, '\n', self.article_text,
                                   flags=re.DOTALL)
        pattern = r'\\keywordsr{(.*?)}+'
        self.article_text = re.sub(pattern, '\n', self.article_text,
                                   flags=re.DOTALL)

    def compile(self):
        """
        Скомпилировать статью.
            1. Выделить необходимую информацию: заголовки, список авторов,
               аннотации и т.д.
            2. Выделить "чистый" текст.
            2. Добавить служебную информацию.
        """
        self.parse()
        m_start = re.search(r'\\markboth', self.text)
        m_end = re.search(r'\\end{thebibliography}', self.text)
        if m_start and m_end:
            self.article_text = self.text[m_start.start():m_end.end()]
            self.remove_russian_abstract()
            self.article_text = \
                r'\input{__init_counters__}' + '\n' + \
                f'\\input{{__to_{self.lang}__}}' + '\n\n' + \
                self.add_content_lines() + \
                fr'\label{{{self.begin_label}}}' + '\n\n' + \
                self.article_text + '\n\n' + \
                fr'\label{{{self.end_label}}}'
            self.update_title()
            for s in self.sections:
                self.update_section(s)
            self.update_image_path()
            with open(self.art_path, 'wt') as f:
                f.write(self.article_text)
        else:
            raise Exception('Incorrect article {}'.format(self.path))

    def as_dict(self):
        return {
            'title': self.title,
            'authors': ['{} {} {}'.format(a['family'], a['name'],
                                          a['patronymic'])
                        for a in self.authors],
            'abstracts': {
                'ru': self.abstracts['ru'],
                'en': self.abstracts['en']
            },
            'УДК': self.udc,
            'MSC2010': self.msc2010,
            'sections': self.sections,
            'text': self.article_text,
            'bibliography': self.bibliography,
        }


class TvimDocument:
    """
    Объектная модель выпуска журнала.
    """
    def __init__(self, config):
        self.config = config
        self.articles = []
        self.verbatim_articles = []
        # parameters
        self.year = self.config['tvim']['year']
        self.number = self.config['tvim']['number']
        self.total_number = self.config['tvim']['total number']
        self.protocol_number = self.config['tvim'].get('protocol number', '???')
        self.protocol_day = self.config['tvim'].get('protocol day', '???')
        self.protocol_month = self.config['tvim'].get('protocol month', '???')
        self.protocol_monthname = self.config['tvim'].get('protocol month name',
                                                          '???')
        self.protocol_year = self.config['tvim'].get('protocol year', '???')
        self.resources = self.config['path']['resources']
        self.page_count = 0

        self.root_path = 'numbers/tvim_{}_{}'.format(self.year, self.number)

    @classmethod
    def from_config(cls, path):
        """
        Parameters
        ----------
            path: str
                Путь к конфигурационному файлу
        """
        with open(path, 'rt') as config_file:
            config = yaml.load(config_file, Loader=yaml.SafeLoader)
        return cls(config)

    art_number = property(lambda self: len(self.articles), None, None)

    def _update_params(self):
        """
        Обновляет основные параметры выпуска журнала, такие как год, номер,
        протокол, даты выхода в свет и др.
        """
        # обновление параметров на русском языке
        params = [
            r'\def\tvimname{Таврический вестник информатики и математики}''\n',
            r'\def\tvimnumber{{№\,{number}\,({total_number})}}''\n'.format(
                number=self.number, total_number=self.total_number),
            r'\def\tvimyear{{{year}}}''\n'.format(year=self.year),
            r'\def\tvimemail{article@tvim{.}info}''\n',
            r'\def\tvimwww{www{.}tvim{.}info}''\n',
            r'\def\protocolnumber{{{}}}'.format(self.protocol_number),
            r'\def\protocolday{{{}}}''\n'.format(self.protocol_day),
            r'\def\protocolmonthname{{{}}}''\n'.format(self.protocol_monthname),
            r'\def\protocolmonth{{{}}}''\n'.format(self.protocol_month),
            r'\def\protocolyear{{{}}}''\n'.format(self.protocol_year),
            r'\def\protocol{№\,\protocolnumber\ от~\protocolday~'
            r'\protocolmonthname~\protocolyear\,г.}''\n',
            r'\def\sign2print{\protocolday.\protocolmonth.\protocolyear}''\n',
            r'\def\print_page_count{{{}}}''\n'.format(
                round(self.page_count * 0.1056), 1),
            r'\def\tvimissn{ISSN\;1729-3901}''\n',
            r'\newlength{\myparindent}''\n',
            r'\newlength{\myinter}''\n'
        ]

        with open('__params__.tex', 'wt') as f:
            f.writelines(params)

        # обновление параметров на английском языке
        params = [
            r'\def\tvimnameen{Taurida Journal of~Computer Science Theory '
            r'and~Mathematics}''\n',
            r'\def\tvimnumberen{{{}}}''\n'.format(self.number),
            r'\def\tvimnumberwithtotalen{{No.\;{}\;({})}}''\n'.format(
                self.number, self.total_number),
            r'\def\tvimyearen{{{}}}''\n'.format(self.year),
            r'\def\tvimemailen{article@tvim{.}info}''\n',
            r'\def\tvimwwwen{www{.}tvim{.}info}''\n',
            r'\def\protocolen{{No.\,? from {}/{}/{}.}}''\n'.format(
                self.protocol_month, self.protocol_day, self.protocol_year),
            r'\def\Profen{Professor}''\n',
            r'\def\Docenten{Associate professor}''\n',
            r'\def\Dfmnen{Doctor of Physico-Mathematical Sciences}''\n',
            r'\def\Dtnen{Doctor of Engineering Sciences}''\n',
            r'\def\Kfmnen{Candidate of Physico-Mathematical Sciences}''\n',
            r'\def\profen{professor}''\n',
            r'\def\docenten{associate professor}''\n',
            r'\def\dfmnen{doctor of Physico-Mathematical Sciences}''\n',
            r'\def\dtnen{doctor of Engineering Sciences}''\n',
            r'\def\kfmnen{candidate of Physico-Mathematical Sciences}''\n',
        ]

        with open('__params_en__.tex', 'wt') as f:
            f.writelines(params)

    @staticmethod
    def _get_referat(article, horzline=True):
        template = '\\tvimRef{{{authors_fio}}}\n' \
                   '{{{authors_iof}}}\n' \
                   '{{{title}}}\n' \
                   '{{{begin_label}}}\n' \
                   '{{{end_label}}}\n' \
                   '{{{udc}}}\n' \
                   '{{{abstract_ru}}}\n' \
                   '{{{keywords_ru}}}\n'
        if horzline:
            template += '\\abstractLine{{0.5cm}}{{0.7cm}}\n\n'

        title = re.sub(r'\\footnote{.*?}', '', article.title['ru'])

        return template.format(authors_fio=article.authors_str,
                               authors_iof=article.authors_str_reverse,
                               title=title,
                               begin_label=article.begin_label,
                               end_label=article.end_label,
                               udc=article.udc,
                               abstract_ru=article.abstracts['ru'],
                               keywords_ru=article.keywords['ru'])

    def _build(self):
        articles_path = os.path.join('articles')
        articles = [
            f for f in os.listdir(articles_path)
            if not f.startswith('.') and not f.startswith('-')
        ]

        self.verbatim_articles = []
        self.articles = []
        author_details = []

        for art in articles:
            art_path = os.path.join(articles_path, art)
            article = None
            if art.startswith('_'):
                article = VerbatimArticle(art_path)
                self.verbatim_articles.append(article)
            elif not art.startswith('-'):
                if art.endswith('_en'):
                    lang = 'eng'
                elif art.endswith('_ukr'):
                    lang = 'ukr'
                else:
                    lang = 'rus'
                article = Article(art_path, lang)
                self.articles.append(article)
            if article:
                article.compile()

        self.articles = sorted(self.articles, key=lambda a: a.authors_str)

        art_file_content = []
        referats = []
        narticles = len(self.articles)
        # collecting verbatim articles
        for i, art in enumerate(self.verbatim_articles):
            art_file_content.append(art.art_path)

        # collecting scientific articles
        for i, art in enumerate(self.articles):
            author_details.extend(art.author_details)
            if art.art_path:
                art_file_content.append(art.art_path)
                referats.append(self._get_referat(art, i < narticles - 1))

        with open('articles.tex', 'wt') as f:
            f.writelines([r'\input{{{}}}''\n'.format(art_path)
                          for art_path in art_file_content])

        author_details = sorted(author_details)
        with open('authors.tex', 'at') as authors_file:
            authors_file.writelines('\n\n'r'\medskip''\n'.join(author_details))

        with open('referats.tex', 'at') as referats_file:
            referats_file.writelines('\n\n'.join(referats))

    def calc_page_count(self, pdf_path):
        self.page_count = 0
        with open(pdf_path, 'rb') as pdf_file:
            reader = PyPDF2.PdfFileReader(pdf_file)
            self.page_count = reader.getNumPages()

    def compile(self):
        print('Компиляция ТВИМ {year} №{number}. '
              'Пожалуйста подождите'.format(year=self.year,
                                            number=self.number))
        # задаем корневую папку и копируем необходимые файлы
        # для компиляции журнала
        if os.path.exists(self.root_path):
            shutil.rmtree(self.root_path)
        shutil.copytree(self.resources, self.root_path)
        shutil.copytree(self.config['path']['articles'],
                        os.path.join(self.root_path, 'articles'))

        cur_dir = os.path.abspath(os.path.curdir)
        try:
            os.chdir(self.root_path)
            main_filename = 'tvim_{}_{}.tex'.format(self.year, self.number)
            os.rename('tvim_main.tex', main_filename)
            self._update_params()
            self._build()
            cmd = ['pdflatex', '-halt-on-error', '-file-line-error',
                   'tvim_{}_{}.tex'.format(self.year, self.number)]
            res = subprocess.run(cmd, stdout=subprocess.PIPE)
            if res.returncode == 0:
                self.calc_page_count('tvim_{}_{}.pdf'.format(self.year,
                                                             self.number))
                self._update_params()
                # второй запуск для создания ссылок и содержания
                subprocess.run(cmd, stdout=subprocess.PIPE)
                print(f'ТВИМ {self.year} {self.number} успешно собран')
                with open(f'tvim_{self.year}_{self.number}.json', 'wt') \
                        as json_file:
                    json.dump(self.as_dict(), json_file, indent=4,
                              sort_keys=False, ensure_ascii=False)
            else:
                print('ОШИБКА: что-то пошло не так. '
                      'Посмотрите, пожалуйста, лог файл.')
        finally:
            os.chdir(cur_dir)

    def as_dict(self):
        return {
            'year': self.year,
            'number': self.number,
            'articles': [a.as_dict() for a in self.articles]
        }


class ReportGenerator:

    def __init__(self, tvim_doc: TvimDocument, config):
        self.config = config
        self.template_path = config['path']['docs']
        self.tvim_doc = tvim_doc
        self.root_path = os.path.join(self.tvim_doc.root_path, 'docs')

    @classmethod
    def from_config(cls, tvim_doc: TvimDocument, path):
        """
        Parameters
        ----------
            tvim_doc: TvimDocument
                Объектная модель журнала
            path: str
                Путь к конфигурационному файлу
        """
        with open(path, 'rt') as config_file:
            config = yaml.load(config_file, Loader=yaml.SafeLoader)
        return cls(tvim_doc, config)

    def build_05_predstavlen(self):
        doc = Document(os.path.join(self.root_path, '05predstavlen.docx'))

        last_digit = self.tvim_doc.art_number % 10
        if last_digit == 1:
            art_word = 'статья'
            pages_word = 'страница'
        elif last_digit in [2, 3, 4]:
            art_word = 'статьи'
            pages_word = 'страницы'
        else:
            art_word = 'статей'
            pages_word = 'страниц'

        new_text = '«Таврический вестник информатики и математики», ' \
                   '{year}, №{number}, {nart} {art_word}, ' \
                   '{pagenum} {pages_word}'.format(
                        year=self.tvim_doc.year,
                        number=self.tvim_doc.number,
                        nart=self.tvim_doc.art_number,
                        art_word=art_word,
                        pagenum=self.tvim_doc.page_count,
                        pages_word=pages_word)

        p = doc.paragraphs[8]
        p.text = ''
        run = p.add_run(new_text + '.')
        run.bold = True
        run.underline = True
        p.paragraph_format.line_spacing = 1.5

        p = doc.paragraphs[21]
        p.text = ''
        run = p.add_run(new_text)
        run.bold = True
        run.underline = True
        new_text_1 = ' научно-технический совет пришёл к заключению, '\
                     'что представленный выпуск научного журнала (сборника) '\
                     'рекомендован к изданию.'
        run = p.add_run(new_text_1)
        run.bold = True
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.alignment = 3

        doc.save(os.path.join(self.root_path, '05predstavlen.docx'))

    def build_06zayavlenie(self):
        doc = Document(os.path.join(self.root_path, '06zayavlenie.docx'))

        p = doc.paragraphs[9]
        p.text = 'Наименование: '
        run = p.add_run('Таврический вестник информатики и математики, '
                        '{}, №{}'.format(self.tvim_doc.year,
                                         self.tvim_doc.number))
        run.bold = True
        run.underline = True

        doc.save(os.path.join(self.template_path, '06zayavlenie.docx'))

    def build_expertiza(self):
        doc = Document(os.path.join(self.template_path,
                                    'Приложение1_Экспертиза публикации.docx'))

        for p_index in [14, 22]:
            p = doc.paragraphs[p_index]
            p.text = ''
            run = p.add_run(
                'научный журнал '
                '«Таврический вестник информатики и математики», '
                '{}, №{}'.format(self.tvim_doc.year, self.tvim_doc.number))
            run.bold = True

        doc.save(os.path.join(self.root_path,
                              'Приложение1_Экспертиза публикации.docx'))

    def build_export_doc(self):
        doc = Document(os.path.join(self.root_path,
                                    'Экспортное заключение.docx'))

        p = doc.paragraphs[14]
        p.text = ''
        p.add_run('Внутривузовская комиссия экспортного контроля рассмотрев ')
        run = p.add_run(
            'научный журнал «Таврический вестник информатики и математики», '
            '{}, №{} '.format(self.tvim_doc.year, self.tvim_doc.number))
        run.bold = True
        p.add_run('подтверждает, что в материале, включающем результаты '
                  'научно-исследовательских, опытно-конструкторских и '
                  'технологических работ, финансируемых государством, '
                  'не содержатся сведения, подпадающие под действие списков '
                  'контролируемых товаров и технологий, и они не могут быть '
                  'использованы для целей создания оружия массового поражения, '
                  'средств его доставки, иных видов вооружения и военной '
                  'техники либо для подготовки и (или) совершения '
                  'террористических актов.')

        p = doc.paragraphs[16]
        p.text = 'Заключение: '
        run = p.add_run('для открытого опубликования подготовленных '
                        'материалов ')
        run.italic = True
        run = p.add_run(
            'научный журнал «Таврический вестник информатики и математики», '
            '{}, №{} '.format(self.tvim_doc.year, self.tvim_doc.number))
        run.bold = True
        run.italic = True
        run = p.add_run('оформление лицензии ФСТЭК России или разрешения '
                        'Комиссии по экспортному контролю '
                        'Российской Федерации не требуется.')
        run.italic = True

        doc.save(os.path.join(self.root_path, 'Экспортное заключение.docx'))

    def build(self):
        print('Создание документов...')
        if os.path.exists(self.root_path):
            shutil.rmtree(self.root_path)
        shutil.copytree(self.template_path, self.root_path)
        self.build_05_predstavlen()
        self.build_06zayavlenie()
        self.build_expertiza()
        self.build_export_doc()
        print('Создание документов успешно завершено')


if __name__ == '__main__':
    argparser = argparse.ArgumentParser(description='TVIM compiler')

    argparser.add_argument('--config', '-C', type=str,
                           default='configs/config.yaml',
                           help='path to config file in YAML format')
    argparser.add_argument('--report', '-R', action='store_true',
                           help='build report')
    args = argparser.parse_args()

    tvim = TvimDocument.from_config(args.config)
    tvim.compile()

    if args.report:
        rep_gen = ReportGenerator.from_config(tvim, args.config)
        rep_gen.build()
